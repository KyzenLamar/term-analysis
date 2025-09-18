# term_analysis_core.py

############################## ВЕРСІЯ З LLM ###################################
import re
import docx
import pdfplumber
import csv
import os

def load_terms_dictionary_csv(file_path):
    """
    Завантаження словника з CSV у форматі:
    approved_term,synonyms,category,wrong_usages,context_examples
    synonyms/wrong_usages/context_examples — через ;
    """
    terms = []
    with open(file_path, encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            terms.append({
                "approved_term": row["approved_term"].strip(),
                "synonyms": [s.strip() for s in row["synonyms"].split(';') if s.strip()],
                "category": row.get("category", "").strip(),
                "wrong_usages": [s.strip() for s in row.get("wrong_usages", "").split(';') if s.strip()],
                "context_examples": [s.strip() for s in row.get("context_examples", "").split(';') if s.strip()],
            })
    return terms

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    # Зберігаємо сторінковість: повертаємо список рядків (розбивка по сторінках)
    return ['\n'.join([p.text for p in doc.paragraphs])]

def extract_text_from_pdf(pdf_path):
    text_by_pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text_by_pages.append(page.extract_text() or "")
    return text_by_pages

def analyze_document_v2(file_path, terms, fuzzy_threshold=88):
    """
    Аналіз документу по новому словнику terms (list of dicts).
    Шукає wrong_usages, synonyms і approved_term.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        doc = docx.Document(file_path)
        text_by_pages = ['\n'.join([p.text for p in doc.paragraphs])]
    elif ext == ".pdf":
        text_by_pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_by_pages.append(page.extract_text() or "")
    else:
        raise ValueError("Unsupported file type: must be .docx or .pdf")

    results = []
    for page_idx, text in enumerate(text_by_pages):
        sentences = re.split(r'[.?!;\n]', text)
        for sentence in sentences:
            s_low = sentence.lower()
            if not sentence.strip():
                continue
            for term in terms:
                # Перевіряємо наявність грубих помилок (wrong_usages)
                for wrong in term["wrong_usages"]:
                    if re.search(r'\b{}\b'.format(re.escape(wrong.lower())), s_low):
                        results.append({
                            "page": page_idx + 1,
                            "type": "wrong_usage",
                            "found_term": wrong,
                            "approved_term": term["approved_term"],
                            "context": sentence.strip(),
                            "category": term["category"],
                        })
                # Перевіряємо approved_term і синоніми
                all_valid = [term["approved_term"]] + term["synonyms"]
                for valid in all_valid:
                    if valid and re.search(r'\b{}\b'.format(re.escape(valid.lower())), s_low):
                        results.append({
                            "page": page_idx + 1,
                            "type": "valid_term",
                            "found_term": valid,
                            "approved_term": term["approved_term"],
                            "context": sentence.strip(),
                            "category": term["category"],
                        })
    return results

def get_unmatched_sentences(file_path, terms, fuzzy_threshold=88):
    """
    Повертає всі речення з документа, які не були знайдені як match у rule-based аналізі.
    (для semantic search)
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        doc = docx.Document(file_path)
        text_by_pages = ['\n'.join([p.text for p in doc.paragraphs])]
    elif ext == ".pdf":
        text_by_pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_by_pages.append(page.extract_text() or "")
    else:
        raise ValueError("Unsupported file type: must be .docx or .pdf")

    found_sentences = set()
    all_sentences = []
    for page_idx, text in enumerate(text_by_pages):
        sentences = re.split(r'[.?!;\n]', text)
        for sentence in sentences:
            s_low = sentence.lower()
            if not sentence.strip():
                continue
            all_sentences.append(sentence.strip())
            for term in terms:
                # Wrong usages
                for wrong in term["wrong_usages"]:
                    if re.search(r'\b{}\b'.format(re.escape(wrong.lower())), s_low):
                        found_sentences.add(sentence.strip())
                # All valid terms
                all_valid = [term["approved_term"]] + term["synonyms"]
                for valid in all_valid:
                    if valid and re.search(r'\b{}\b'.format(re.escape(valid.lower())), s_low):
                        found_sentences.add(sentence.strip())
    # Повертаємо лише ті, що не були знайдені rule-based
    unmatched = [s for s in all_sentences if s not in found_sentences]
    return unmatched

def export_results_to_csv(results, out_path):
    """
    Зберігає результати у .csv (з новими полями)
    """
    fieldnames = ["page", "type", "found_term", "approved_term", "context", "category"]
    with open(out_path, 'w', encoding='utf-8', newline='') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for row in results:
            writer.writerow(row)

# ==== Як використовувати ====
if __name__ == "__main__":
    synonyms_path = "synonyms_output.csv"
    file_path = "example.docx"  # або .pdf
    out_path = "report.csv"

    terms = load_terms_dictionary_csv(synonyms_path)
    results = analyze_document_v2(file_path, terms)
    export_results_to_csv(results, out_path)
    print(f"Done! Results saved to {out_path}")

